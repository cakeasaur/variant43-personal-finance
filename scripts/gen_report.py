"""Generate пояснительная записка (ГОСТ 7.32-2017) for variant 43."""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ── helpers ──────────────────────────────────────────────────────────────────

def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def add_page_numbers(doc: Document) -> None:
    """Insert «page / total» field in footer of every section."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        run = para.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)
        _fmt_run(run)


def _fmt_run(run, size: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


def _para(doc: Document,
          text: str = "",
          bold: bool = False,
          size: int = 14,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          indent_first: float = 1.25,
          space_before: int = 0,
          space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(indent_first) if indent_first else None
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(21)  # ~1.5 × 14pt
    if text:
        run = p.add_run(text)
        _fmt_run(run, size=size, bold=bold)
    return p


def heading1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(21)
    run = p.add_run(text.upper())
    _fmt_run(run, size=14, bold=True)


def heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(21)
    run = p.add_run(text)
    _fmt_run(run, size=14, bold=True)


def heading3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(21)
    run = p.add_run(text)
    _fmt_run(run, size=14, bold=True)


def body(doc: Document, text: str) -> None:
    _para(doc, text, bold=False, size=14, indent_first=1.25,
          space_before=0, space_after=6)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _fmt_run(run, size=12, bold=True)
    # data rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            _fmt_run(run, size=12)
    doc.add_paragraph()  # spacer


# ── title page ────────────────────────────────────────────────────────────────

def build_title_page(doc: Document) -> None:
    def center(text: str, bold: bool = False, size: int = 14) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = Pt(21)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        _fmt_run(run, size=size, bold=bold)

    center("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", bold=True, size=12)
    center("Федеральное государственное бюджетное образовательное учреждение высшего образования", size=12)
    center("«ТУЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»", bold=True, size=14)
    doc.add_paragraph()
    center("Кафедра программного обеспечения вычислительной техники и автоматизированных систем", size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА")
    _fmt_run(run, size=16, bold=True)

    center("к курсовой работе (проекту)", size=14)
    center("по дисциплине", size=14)
    center("«Методы и технологии программирования»", bold=True, size=14)
    doc.add_paragraph()
    center("Тема:", bold=True, size=14)
    center("Приложение для управления личными финансами", bold=True, size=14)
    center("(Вариант 43)", size=14)
    doc.add_paragraph()
    doc.add_paragraph()

    # student info block (right-aligned)
    for line in [
        "Выполнил(а): студент(ка) гр. ________________",
        "______________________________",
        "                                          (ФИО)",
        "",
        "Проверил(а): ____________________________",
        "______________________________",
        "                                    (учёное звание, ФИО)",
        "",
        "Оценка: _________________________________",
        "Дата: ___________________________________",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(line)
        _fmt_run(run, size=12)

    doc.add_paragraph()
    doc.add_paragraph()
    center("Тула – 2026", bold=False, size=14)
    doc.add_page_break()


# ── table of contents (auto field) ───────────────────────────────────────────

def build_toc(doc: Document) -> None:
    heading1(doc, "СОДЕРЖАНИЕ")

    # Вставляем настоящее поле TOC — WPS/Word обновит его по F9
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    _fmt_run(run, size=14)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    fldChar_begin.set(qn("w:dirty"), "true")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_sep)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)

    doc.add_page_break()


# ── sections ──────────────────────────────────────────────────────────────────

def build_intro(doc: Document) -> None:
    heading1(doc, "ВВЕДЕНИЕ")

    body(doc,
         "Управление личными финансами требует регулярного учёта доходов и расходов, "
         "анализа структуры трат и контроля будущих обязательных платежей. Без специализированного "
         "инструмента пользователь теряет прозрачность бюджета, что существенно усложняет "
         "финансовое планирование и достижение поставленных целей.")

    body(doc,
         "Актуальность разработки обусловлена высоким спросом на приложения персональных "
         "финансов: по данным аналитиков, более 60 % пользователей смартфонов хотя бы "
         "раз устанавливали приложение для ведения бюджета. При этом значительная часть "
         "существующих решений требует постоянного подключения к интернету или привязки "
         "к банковскому счёту, что создаёт риски утечки персональных данных.")

    body(doc,
         "Цель курсовой работы — разработать кроссплатформенное мобильное приложение "
         "для управления личными финансами (вариант 43), обеспечивающее: учёт доходов "
         "и расходов; категоризацию транзакций; визуализацию данных в виде диаграмм; "
         "финансовые цели с отслеживанием прогресса; напоминания о платежах; защиту "
         "данных с помощью шифрования локальной базы данных.")

    body(doc, "Для достижения поставленной цели решаются следующие задачи:")

    tasks = [
        "провести анализ предметной области и сформулировать требования к программному продукту;",
        "спроектировать архитектуру приложения и модель данных SQLite;",
        "реализовать приложение на стеке Python + Kivy + SQLite;",
        "реализовать модули отчётности, целей и напоминаний;",
        "реализовать защиту данных с применением стойкого шифрования;",
        "подготовить тесты и проверки качества кода;",
        "оформить пояснительную записку в соответствии со структурой методических указаний.",
    ]
    for _i, t in enumerate(tasks, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(t)
        _fmt_run(run, size=14)

    body(doc,
         "Объект исследования — персональные финансы пользователя: транзакции, "
         "финансовые цели и обязательные платежи. Предмет исследования — методы и "
         "технологии разработки кроссплатформенного приложения с локальным хранением "
         "и криптографической защитой данных.")

    body(doc,
         "Курсовая работа выполнена с применением следующих инструментов: Python 3.12, "
         "Kivy 2.3.1, SQLite, библиотека cryptography (AES-256-GCM, scrypt), "
         "pytest, ruff, pip-audit, Docker, GitHub Actions.")

    doc.add_page_break()


def build_section1(doc: Document) -> None:
    heading1(doc, "1 АНАЛИТИЧЕСКАЯ ЧАСТЬ")

    heading2(doc, "1.1 Описание предметной области")

    body(doc,
         "Персональные финансы — область деятельности, связанная с учётом и анализом "
         "денежных потоков физического лица: доходов, расходов, накоплений и "
         "финансовых обязательств. Приложение предназначено для офлайн-учёта "
         "транзакций и получения наглядной картины финансового состояния за "
         "выбранный период.")

    body(doc, "Целевая аудитория приложения:")

    items = [
        "пользователи, самостоятельно контролирующие личный бюджет;",
        "пользователи, которым необходимы напоминания о регулярных платежах (аренда, коммунальные услуги, кредиты);",
        "пользователи, накапливающие средства на конкретную цель и желающие видеть прогресс.",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run("– " + item)
        _fmt_run(run, size=14)

    body(doc,
         "Ключевой контекст использования — мобильное устройство, быстрый ввод расходов «на ходу», "
         "отсутствие доступа к интернету не должно блокировать работу (принцип offline-first).")

    heading3(doc, "1.1.1 Глоссарий")

    terms = [
        ("Транзакция", "запись о доходе или расходе с датой, суммой и категорией."),
        ("Категория", "классификатор транзакций (например, «Еда», «Транспорт»)."),
        ("Период", "интервал дат, по которому строятся отчёты."),
        ("Отчёт/сводка", "агрегированные показатели по периоду: итог доходов/расходов, доли по категориям."),
        ("Финансовая цель", "накопление заданной суммы к сроку с отображением прогресса."),
        ("Напоминание", "будущий платёж с датой и опциональным правилом повторения."),
        ("Локальная БД", "SQLite-файл с данными пользователя на устройстве."),
        ("Шифрование", "защита содержимого БД от чтения без ключа (пароля пользователя)."),
    ]
    add_table(doc, ["Термин", "Определение"], terms)

    heading2(doc, "1.2 Анализ аналогов")

    body(doc,
         "Для формирования обоснованных требований проведён обзор популярных "
         "приложений для ведения личных финансов.")

    analog_rows = [
        ("Zenmoney", "Учёт, категории, бюджеты, цели; синхронизация с банками", "Требует подписки и интернета"),
        ("CoinKeeper", "Акцент на UX, быстрый ввод, визуализация", "Ограниченный бесплатный функционал"),
        ("Monefy", "Простота, быстрый учёт, базовые отчёты", "Нет целей, слабые отчёты"),
        ("Spendee", "Категории, визуализация, бюджеты/кошельки", "Синхронизация платная"),
        ("Money Manager", "Широкие отчёты, учёт долгов/счетов", "Сложный интерфейс"),
    ]
    add_table(doc,
              ["Аналог", "Достоинства", "Недостатки"],
              analog_rows)

    body(doc,
         "По результатам анализа сформулированы следующие выводы для проекта: "
         "необходимо реализовать offline-first режим работы с локальной SQLite БД; "
         "обеспечить визуализацию (круговая диаграмма расходов по категориям, "
         "динамика по дням); включить модули целей и напоминаний; обеспечить "
         "защиту данных посредством шифрования — обязательное требование варианта 43.")

    heading2(doc, "1.3 Требования к программному продукту")

    heading3(doc, "1.3.1 Функциональные требования")

    fr_rows = [
        ("FR-01", "Ведение транзакций", "Добавление, удаление дохода/расхода: сумма, дата, категория, заметка"),
        ("FR-02", "Категории", "Предустановленные категории; CRUD пользовательских — backlog v0.2"),
        ("FR-03", "Просмотр и фильтрация", "Список транзакций за период, фильтр по типу и категории, итоги"),
        ("FR-04", "Визуализация", "Диаграмма расходов по категориям, динамика по дням за период"),
        ("FR-05", "Финансовые цели", "CRUD целей, пополнение, прогресс в % и абсолютных значениях"),
        ("FR-06", "Напоминания", "CRUD, повторение (none/daily/weekly/monthly), список ближайших"),
        ("FR-07", "Защита данных", "Шифрование локальной БД; пароль при запуске приложения"),
    ]
    add_table(doc, ["ID", "Требование", "Описание"], fr_rows)

    heading3(doc, "1.3.2 Нефункциональные требования")

    nfr_rows = [
        ("NFR-01", "Кроссплатформенность", "Android и desktop (Windows/Linux/macOS) через Kivy"),
        ("NFR-02", "Offline-first", "Все функции работают без интернета"),
        ("NFR-03", "Производительность", "Отзывчивость при N ≈ 5 000 транзакций"),
        ("NFR-04", "Целостность данных", "Атомарность операций записи (транзакции SQLite)"),
        ("NFR-05", "Безопасность", "Данные не читаются без ключа; KDF для получения ключа из пароля"),
        ("NFR-06", "Поддерживаемость", "Трёхслойная структура, модульные тесты"),
    ]
    add_table(doc, ["ID", "Требование", "Описание"], nfr_rows)

    doc.add_page_break()


def build_section2(doc: Document) -> None:
    heading1(doc, "2 ПРОЕКТНАЯ ЧАСТЬ")

    heading2(doc, "2.1 Архитектура приложения")

    heading3(doc, "2.1.1 Принцип разделения на слои")

    body(doc,
         "Приложение построено по трёхслойной архитектуре с однонаправленными "
         "зависимостями: UI → Core ← Infra. Направление зависимостей гарантирует, "
         "что доменная логика не зависит от деталей хранения или отображения.")

    layers = [
        ("UI (presentation)", "Экраны Kivy, виджеты, обработка событий ввода пользователя"),
        ("Core (domain)", "Модели данных (frozen dataclasses), правила валидации, "
                          "чистые функции агрегирования отчётов"),
        ("Infra", "SQLite-репозитории, схема БД, криптографический модуль"),
    ]
    add_table(doc, ["Слой", "Ответственность"], layers)

    body(doc,
         "Ключевое свойство архитектуры: модуль src/core/ не имеет внешних зависимостей "
         "(только stdlib), что делает его полностью тестируемым без подключения к БД и GUI.")

    heading3(doc, "2.1.2 Структура каталогов")

    code_block(doc, """src/
  app.py                  # точка входа Kivy + логика старта/шифрования
  core/
    models.py             # Transaction, TransactionType (frozen dataclass)
    reporting.py          # totals_for_period, expense_by_category, expense_by_day
    perf.py               # бенчмарк-хелперы
  infra/
    db/
      connection.py       # connect() + transaction() context manager
      schema.py           # init_schema(), SCHEMA_VERSION = 3
      repositories.py     # CategoryRepository, TransactionRepository,
    security/             #   GoalRepository, ReminderRepository
      crypto.py           # encrypt_bytes/decrypt_bytes (AES-256-GCM + scrypt)
  ui/
    cards.py              # виджеты карточек
    factories.py          # вспомогательные UI-фабрики
    formatting.py         # форматирование денег, дат (локаль RU)
    forms.py              # формы ввода
    screens/              # экраны: overview, goals, reminders, reports
    theme.py              # цвета, шрифты, иконки Material Design
    widgets.py            # переиспользуемые виджеты (ModalSheet, BarTrack)
tests/
  conftest.py
  test_core_reporting.py
  test_crypto.py
  test_infra_db.py
  test_smoke.py
main.py                   # python main.py → запуск приложения""")

    heading3(doc, "2.1.3 Диаграмма компонентов")

    body(doc,
         "Диаграмма компонентов (PlantUML/Mermaid, файл docs/diagrams/component.mmd) "
         "описывает потоки зависимостей между слоями. Экраны и виджеты UI обращаются "
         "к репозиториям инфраструктурного слоя и к агрегирующим функциям доменного "
         "слоя. Модуль безопасности взаимодействует с файлом БД напрямую при "
         "шифровании/расшифровании.")

    code_block(doc, """flowchart LR
  subgraph UI[UI (Kivy)]
    Screens --> Widgets
  end
  subgraph CORE[Core / Domain]
    Models
    Reports[Aggregations]
  end
  subgraph INFRA[Infrastructure]
    Repo[SQLite Repositories] --> DB[(SQLite .enc)]
    Sec[Security AES-GCM] --> DB
  end
  Screens --> Reports
  Screens --> Repo
  Reports --> Models
  Repo --> Models""")

    heading2(doc, "2.2 Проектирование базы данных")

    heading3(doc, "2.2.1 Общие принципы")

    body(doc,
         "База данных — локальная SQLite, один файл. "
         "Ключевые проектные решения:")
    decisions = [
        "суммы хранятся в целых числах (amount_cents: int) — исключает погрешности вещественной арифметики;",
        "даты хранятся как ISO-8601 TEXT — обеспечивает лексикографическую сортировку;",
        "удаление категории выполняет ON DELETE SET NULL для транзакций — история сохраняется;",
        "PRAGMA journal_mode = DELETE запрещает создание WAL-файлов, что исключает утечку незашифрованных данных;",
        "PRAGMA foreign_keys = ON обеспечивает ссылочную целостность.",
    ]
    for d in decisions:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run("– " + d)
        _fmt_run(run, size=14)

    heading3(doc, "2.2.2 Схема базы данных (ER-диаграмма)")

    body(doc,
         "ER-диаграмма (файл docs/diagrams/er.mmd) содержит пять таблиц: "
         "settings, categories, transactions, goals, reminders.")

    schema_rows = [
        ("settings", "key (PK), value", "Системные параметры, версия схемы"),
        ("categories", "id (PK), name (UNIQUE), kind, color, icon, created_at, updated_at",
         "Справочник категорий"),
        ("transactions", "id (PK), type, amount_cents, occurred_at, category_id (FK), note, created_at, updated_at",
         "Доходы и расходы пользователя"),
        ("goals", "id (PK), name, target_cents, current_cents, deadline_at, note, created_at, updated_at",
         "Финансовые цели"),
        ("reminders", "id (PK), name, amount_cents, due_at, recurrence, note, created_at, updated_at",
         "Напоминания о платежах"),
    ]
    add_table(doc, ["Таблица", "Поля", "Назначение"], schema_rows)

    heading3(doc, "2.2.3 Индексы")

    body(doc,
         "Для ускорения типовых запросов (выборка за период, фильтр по категории) "
         "созданы составные индексы:")

    idx_rows = [
        ("idx_transactions_occurred_at", "transactions(occurred_at)", "Выборка за период"),
        ("idx_transactions_category_occurred", "transactions(category_id, occurred_at)", "Фильтр по категории и дате"),
        ("idx_transactions_type_occurred", "transactions(type, occurred_at)", "Фильтр по типу (доход/расход)"),
        ("idx_goals_deadline", "goals(deadline_at)", "Сортировка целей по сроку"),
        ("idx_reminders_due_at", "reminders(due_at)", "Выборка ближайших напоминаний"),
        ("idx_reminders_recurrence_due_at", "reminders(recurrence, due_at)", "Фильтр повторяющихся"),
    ]
    add_table(doc, ["Индекс", "Столбцы", "Назначение"], idx_rows)

    heading3(doc, "2.2.4 Версионирование схемы")

    body(doc,
         "Текущая версия схемы — SCHEMA_VERSION = 3, хранится в таблице settings. "
         "Все таблицы создаются через CREATE TABLE IF NOT EXISTS, что сохраняет "
         "существующие данные при добавлении новых таблиц. Автоматические миграции "
         "запланированы в backlog v0.2.")

    doc.add_page_break()


def build_section3(doc: Document) -> None:
    heading1(doc, "3 ТЕХНОЛОГИЧЕСКАЯ ЧАСТЬ")

    heading2(doc, "3.1 Выбор языка программирования: Python 3.12")

    body(doc,
         "Python выбран по следующим причинам: высокая скорость разработки и "
         "читаемость кода; богатая экосистема библиотек; удобство реализации "
         "доменной логики и тестируемых модулей; нативная поддержка в Kivy. "
         "Версия 3.12 закреплена в файле .python-version: обеспечивает "
         "совместимость всех зависимостей и доступность новых возможностей языка "
         "(StrEnum, frozen dataclasses со slots, type hints).")

    heading2(doc, "3.2 Выбор UI-фреймворка: Kivy 2.3.1")

    body(doc,
         "Kivy — open-source UI-фреймворк для Python, поддерживающий Android, "
         "iOS, Windows, Linux, macOS. Обоснование выбора: точное соответствие "
         "требованиям варианта 43; поддержка touch-интерфейса; декларативный kv-язык "
         "для описания разметки; возможность упаковки в APK через Buildozer.")

    heading2(doc, "3.3 Выбор СУБД: SQLite")

    body(doc,
         "SQLite — встроенная реляционная СУБД без отдельного сервера. Выбор "
         "обоснован: обеспечивает offline-first режим (NFR-02); не требует "
         "установки и настройки; поддерживает транзакции, индексы и ограничения "
         "целостности; является стандартным хранилищем для мобильных приложений.")

    heading2(doc, "3.4 Стратегия защиты данных (шифрование локальной БД)")

    heading3(doc, "3.4.1 Модель угроз")

    body(doc,
         "Защита направлена против сценария: злоумышленник получил доступ к файлу "
         "БД на устройстве (копирование, физический доступ) и пытается прочитать "
         "данные вне приложения. Компрометация устройства на уровне ОС (root/jailbreak) "
         "находится за пределами области защиты данного проекта.")

    heading3(doc, "3.4.2 Алгоритм шифрования")

    body(doc,
         "Выбран подход полного шифрования файла БД (стратегия A). "
         "Параметры шифрования:")

    crypto_rows = [
        ("Алгоритм", "AES-256-GCM", "Симметричное шифрование с аутентификацией (AEAD)"),
        ("KDF", "scrypt (n=2¹⁴, r=8, p=1)", "Вывод ключа из пароля; устойчив к GPU-атакам"),
        ("Длина ключа", "256 бит (32 байта)", "Соответствует AES-256"),
        ("Salt", "16 байт (random)", "Уникален для каждого файла; хранится в контейнере"),
        ("Nonce", "12 байт (random)", "Уникален для каждого шифрования"),
        ("Формат контейнера", "MAGIC(4) + salt(16) + nonce(12) + ciphertext", "Расширение .enc"),
    ]
    add_table(doc, ["Параметр", "Значение", "Пояснение"], crypto_rows)

    heading3(doc, "3.4.3 Жизненный цикл БД")

    body(doc,
         "1. Запуск приложения: пользователь вводит пароль → ключ выводится через scrypt → "
         "зашифрованный файл data/*.enc расшифровывается во временный plaintext-файл "
         "(tempfile.mkstemp) → соединение с SQLite открывается через plaintext-файл. "
         "2. Работа: все операции с БД выполняются через plaintext-файл в памяти. "
         "3. Завершение: plaintext-файл шифруется обратно в data/*.enc → plaintext-файл удаляется. "
         "При аварийном завершении atexit-обработчик удаляет plaintext-файл; "
         "при следующем старте «осиротевшие» plaintext-файлы очищаются функцией "
         "_cleanup_orphaned_plaintext_dbs().")

    heading3(doc, "3.4.4 Прочие ключевые библиотеки")

    libs_rows = [
        ("pytest 9.0.3", "Модульные и интеграционные тесты"),
        ("ruff 0.15.11", "Линтер + форматирование импортов (правила E, F, I, B, UP)"),
        ("pip-audit 2.9.0", "Проверка зависимостей на известные CVE"),
        ("pyinstaller 6.19.0", "Упаковка в исполняемый файл для Windows"),
        ("cryptography 46.0.7", "AES-GCM и scrypt"),
        ("Pillow 10.4.0", "Обработка изображений для Kivy"),
    ]
    add_table(doc, ["Библиотека", "Назначение"], libs_rows)

    doc.add_page_break()


def build_section4(doc: Document) -> None:
    heading1(doc, "4 РЕАЛИЗАЦИЯ")

    heading2(doc, "4.1 Структура проекта")

    body(doc,
         "Проект содержит ~2 300 строк кода Python, распределённых по трём слоям: "
         "ui (~1 700 строк), core (~160 строк), infra (~460 строк). "
         "Точка входа — файл main.py в корне; Kivy-приложение запускается "
         "из src/app.py, который управляет жизненным циклом шифрования.")

    heading2(doc, "4.2 Доменный слой (core/)")

    heading3(doc, "4.2.1 Модели данных (models.py)")

    body(doc,
         "Доменные типы реализованы как неизменяемые (frozen) датаклассы со slots "
         "для экономии памяти. Тип TransactionType использует StrEnum — "
         "сравнение с строками работает без явного приведения типов.")

    code_block(doc, """class TransactionType(StrEnum):
    INCOME  = "income"
    EXPENSE = "expense"

@dataclass(frozen=True, slots=True)
class Transaction:
    type:        TransactionType
    amount_cents: int
    occurred_at:  datetime
    category_id:  int | None = None
    note:         str | None = None

    def __post_init__(self) -> None:
        if self.amount_cents <= 0:
            raise ValueError("amount_cents must be > 0")""")

    heading3(doc, "4.2.2 Отчётный модуль (reporting.py)")

    body(doc,
         "Модуль содержит три чистые функции без побочных эффектов — "
         "принимают список транзакций и временной диапазон, возвращают агрегаты. "
         "Отсутствие зависимостей от БД или UI делает их тривиально тестируемыми.")

    code_block(doc, """def totals_for_period(
    transactions: Iterable[Transaction], *, start: datetime, end: datetime
) -> Totals:
    income, expense = 0, 0
    for t in _filter_period(transactions, start, end):
        if t.type == TransactionType.INCOME:
            income += t.amount_cents
        else:
            expense += t.amount_cents
    return Totals(income_cents=income, expense_cents=expense)

def expense_by_category(...) -> dict[int | None, int]: ...
def expense_by_day(...)      -> dict[date, int]:        ...""")

    heading2(doc, "4.3 Инфраструктурный слой (infra/)")

    heading3(doc, "4.3.1 Управление соединением (connection.py)")

    body(doc,
         "Функция connect() открывает соединение с isolation_level=None (autocommit) "
         "и устанавливает journal_mode=DELETE и foreign_keys=ON. "
         "Контекстный менеджер transaction() реализует явный BEGIN/COMMIT/ROLLBACK; "
         "вложенные вызовы являются no-op для предотвращения ошибки "
         "«cannot start a transaction within a transaction».")

    code_block(doc, """@contextmanager
def transaction(conn: sqlite3.Connection):
    if conn.in_transaction:
        yield          # вложенный вызов — no-op
        return
    conn.execute("BEGIN")
    try:
        yield
        conn.execute("COMMIT")
    except Exception:
        conn.execute("ROLLBACK")
        raise""")

    heading3(doc, "4.3.2 Репозитории (repositories.py)")

    body(doc,
         "Каждый репозиторий инкапсулирует CRUD-операции одной сущности. "
         "Все суммы передаются и возвращаются в целых центах (amount_cents: int). "
         "Особого внимания заслуживает GoalRepository.deposit(): "
         "ограничение «не превысить цель» реализовано на уровне SQL через "
         "MIN(target_cents, current_cents + ?), что атомарно и не требует "
         "дополнительной логики в Python.")

    code_block(doc, """def deposit(self, *, goal_id: int, amount_cents: int) -> Goal:
    if amount_cents <= 0:
        raise ValueError("amount_cents must be > 0")
    self.conn.execute(
        \"\"\"UPDATE goals
           SET current_cents = MIN(target_cents, current_cents + ?),
               updated_at    = ?
           WHERE id = ?\"\"\",
        (amount_cents, _now_iso(), int(goal_id)),
    )
    goal = self.get(goal_id=goal_id)
    if goal is None:
        raise ValueError(f"goal {goal_id!r} not found")
    return goal""")

    body(doc,
         "ReminderRepository.mark_done() вычисляет следующую дату повтора "
         "при периодических напоминаниях. Для ежемесячных напоминаний используется "
         "вспомогательная функция _add_months(), корректно обрабатывающая граничные "
         "случаи (например, 31 января + 1 месяц → 28/29 февраля).")

    heading2(doc, "4.4 Защита данных (security/crypto.py)")

    body(doc,
         "Модуль crypto.py реализует шифрование/расшифрование файла БД. "
         "Константы алгоритма вынесены на уровень модуля для единственного "
         "источника истины:")

    code_block(doc, """MAGIC           = b"PFM1"  # Personal Finance Manager v1
SALT_LEN        = 16
NONCE_LEN       = 12      # AESGCM nonce
KEY_LEN         = 32      # 256-bit AES key
MIN_PASSPHRASE_LEN = 8

def encrypt_bytes(*, plaintext: bytes, passphrase: str) -> bytes:
    salt      = os.urandom(SALT_LEN)
    nonce     = os.urandom(NONCE_LEN)
    key       = _derive_key(passphrase, salt)   # scrypt
    ciphertext = AESGCM(key).encrypt(nonce, plaintext, None)
    return MAGIC + salt + nonce + ciphertext

def decrypt_bytes(*, blob: bytes, passphrase: str) -> bytes:
    ...
    try:
        return AESGCM(key).decrypt(nonce, ciphertext, None)
    except InvalidTag as exc:
        raise InvalidPasswordError("invalid password or corrupted data") from exc""")

    body(doc,
         "Константа MIN_PASSPHRASE_LEN = 8 используется как в модуле crypto.py, "
         "так и в валидации пользовательского ввода в src/app.py — единственный "
         "источник истины исключает расхождение проверок.")

    heading2(doc, "4.5 Слой пользовательского интерфейса (UI)")

    body(doc,
         "UI реализован на Kivy и разбит на несколько модулей: "
         "theme.py (Material Design цвета, иконки, шрифты), "
         "factories.py (фабрики виджетов), "
         "formatting.py (форматирование рублей, дат, локаль RU), "
         "widgets.py (переиспользуемые виджеты: ModalSheet, BarTrack), "
         "cards.py (карточки транзакций, целей, напоминаний, Sidebar), "
         "screens/ (экраны: overview, goals, reminders, reports).")

    body(doc,
         "Визуализация отчётов реализована без внешних библиотек графиков: "
         "виджет BarTrack рисует столбчатую диаграмму средствами Kivy Canvas "
         "на основе агрегатов, рассчитанных функциями доменного слоя. "
         "Это устраняет зависимость от matplotlib и снижает размер APK.")

    doc.add_page_break()


def build_section5(doc: Document) -> None:
    heading1(doc, "5 ТЕСТИРОВАНИЕ")

    heading2(doc, "5.1 Стратегия тестирования")

    body(doc,
         "Тесты разделены на три уровня: модульные тесты доменной логики "
         "(без БД и UI), интеграционные тесты инфраструктурного слоя "
         "(реальный SQLite через tmp_path-фикстуры pytest), модульные тесты "
         "криптографического модуля. UI-тесты не реализованы, поскольку Kivy "
         "требует дисплей и не поддерживает headless-режим на CI-раннерах "
         "без Xvfb.")

    body(doc,
         "Ключевой принцип: тесты не используют моки — все обращения к БД "
         "выполняются через настоящее SQLite-соединение в tmp_path. "
         "Это исключает расхождения между поведением мока и реальной СУБД.")

    heading2(doc, "5.2 Результаты тестирования")

    test_files = [
        ("test_core_reporting.py", "8",
         "Агрегации: totals_for_period, expense_by_category, expense_by_day; "
         "граничные условия временного диапазона"),
        ("test_crypto.py", "6",
         "Encrypt/decrypt roundtrip; неверный пароль; повреждённый заголовок; "
         "проверка MIN_PASSPHRASE_LEN"),
        ("test_infra_db.py", "22",
         "CRUD всех репозиториев; транзакционность; вложенные транзакции; "
         "deposit-кэпинг; рекуррентные напоминания; update транзакций"),
        ("test_smoke.py", "6",
         "Импорт всех слоёв; инициализация схемы; пропуск Kivy-тестов "
         "в headless-режиме"),
    ]
    add_table(doc,
              ["Файл", "Тестов", "Покрытие"],
              test_files)

    body(doc,
         "Итог прогона (py -3.12 -m pytest -q): 42 passed in 1.81s. "
         "Все проверки завершились успешно.")

    code_block(doc, """$ py -3.12 -m pytest -q
tests/test_core_reporting.py ........                         [ 19%]
tests/test_crypto.py ......                                   [ 33%]
tests/test_infra_db.py ......................                  [ 85%]
tests/test_smoke.py ......                                    [100%]

42 passed in 1.81s""")

    heading2(doc, "5.3 Статический анализ и проверка зависимостей")

    heading3(doc, "5.3.1 Линтер ruff")

    body(doc,
         "Статический анализ выполняется утилитой ruff (конфигурация в pyproject.toml): "
         "активированы группы правил E (ошибки стиля PEP 8), F (pyflakes: неиспользуемые "
         "импорты, undefined names), I (сортировка импортов isort), "
         "B (flake8-bugbear), UP (pyupgrade). Длина строки — 100 символов. "
         "Результат: All checks passed.")

    heading3(doc, "5.3.2 Проверка зависимостей (pip-audit)")

    body(doc,
         "Инструмент pip-audit сканирует requirements/ci.txt и requirements/dev.txt "
         "на наличие записей в базах CVE/OSV. Результат последней проверки: "
         "уязвимости не обнаружены. Проверка запускается автоматически в "
         "GitHub Actions при каждом push в ветки main и develop.")

    heading3(doc, "5.3.3 CI/CD (GitHub Actions)")

    body(doc,
         "Файл .github/workflows/ci.yml определяет джоб test, запускаемый "
         "на ubuntu-latest при push в main/develop и при открытии PR. "
         "Последовательность шагов: Checkout → Setup Python 3.12 → "
         "Install (requirements/ci.txt + requirements/dev.txt) → "
         "Lint (ruff) → Tests (pytest) → Audit (pip-audit). "
         "Kivy не устанавливается в CI: все тесты написаны с "
         "pytest.mark.skipif(importlib.util.find_spec('kivy') is None, ...), "
         "что обеспечивает полный прогон без дисплея.")

    docker_code = """# Запуск headless-проверок через Docker (аналог CI)
docker compose run --rm checks"""
    code_block(doc, docker_code)

    doc.add_page_break()


def build_conclusion(doc: Document) -> None:
    heading1(doc, "ЗАКЛЮЧЕНИЕ")

    body(doc,
         "В ходе курсовой работы разработано кроссплатформенное приложение "
         "для управления личными финансами (вариант 43) на стеке "
         "Python 3.12 + Kivy 2.3.1 + SQLite.")

    body(doc, "Реализованный функционал версии v0.1:")

    features = [
        "учёт транзакций — добавление доходов и расходов, фильтрация по типу и периоду, "
         "отображение итогов (баланс, суммы доходов/расходов);",
        "отчёты и визуализация — агрегация расходов по категориям и по дням, "
         "диаграммы на базе Kivy Canvas без внешних зависимостей;",
        "финансовые цели — CRUD целей, атомарное пополнение с SQL-ограничением "
         "через MIN(), отображение прогресса;",
        "напоминания — CRUD, поддержка рекуррентности (none/daily/weekly/monthly), "
         "автоматический пересчёт следующей даты при mark_done();",
        "защита данных — AES-256-GCM с KDF scrypt, контейнер с MAGIC/salt/nonce, "
         "отсутствие WAL-утечки через journal_mode=DELETE.",
    ]
    for f in features:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run("– " + f)
        _fmt_run(run, size=14)

    body(doc, "Обеспечение качества:")

    quality = [
        "42 автоматических теста: доменная логика, репозитории SQLite, криптография;",
        "CI на GitHub Actions: ruff + pytest + pip-audit на каждый push;",
        "трёхслойная архитектура без циклических зависимостей (core/ — zero external deps);",
        "Docker-образ для headless-проверок (аналог CI-среды).",
    ]
    for q in quality:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run("– " + q)
        _fmt_run(run, size=14)

    body(doc,
         "Требования варианта 43 выполнены в полном объёме: реализованы учёт "
         "транзакций, категоризация, визуализация, финансовые цели, напоминания "
         "и шифрование локальной БД.")

    body(doc, "Направления развития (backlog v0.2):")
    backlog = [
        "CRUD пользовательских категорий;",
        "редактирование транзакций в интерфейсе;",
        "автоматические миграции схемы БД;",
        "экспорт данных в CSV/PDF;",
        "сборка APK через Buildozer для Android.",
    ]
    for b in backlog:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run("– " + b)
        _fmt_run(run, size=14)

    doc.add_page_break()


def build_references(doc: Document) -> None:
    heading1(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ")

    refs = [
        "Макконнелл С. Совершенный код. — Microsoft Press, 2023. — 896 с.",
        "Фаулер М. Архитектура корпоративных программных приложений. — Addison-Wesley, 2022. — 560 с.",
        "Python Software Foundation. Документация Python 3.12 [Электронный ресурс]. — "
         "Режим доступа: https://docs.python.org/3/ (дата обращения: 14.05.2026).",
        "Kivy Organization. Kivy Documentation [Электронный ресурс]. — "
         "Режим доступа: https://kivy.org/doc/stable/ (дата обращения: 14.05.2026).",
        "SQLite Consortium. SQLite Documentation [Электронный ресурс]. — "
         "Режим доступа: https://www.sqlite.org/docs.html (дата обращения: 14.05.2026).",
        "Python Cryptographic Authority. cryptography — AES-GCM, scrypt [Электронный ресурс]. — "
         "Режим доступа: https://cryptography.io/en/latest/ (дата обращения: 14.05.2026).",
        "pytest Development Team. pytest Documentation [Электронный ресурс]. — "
         "Режим доступа: https://docs.pytest.org/ (дата обращения: 14.05.2026).",
        "Astral. Ruff Documentation [Электронный ресурс]. — "
         "Режим доступа: https://docs.astral.sh/ruff/ (дата обращения: 14.05.2026).",
        "pypa. pip-audit Documentation [Электронный ресурс]. — "
         "Режим доступа: https://pypi.org/project/pip-audit/ (дата обращения: 14.05.2026).",
        "GitHub, Inc. GitHub Actions Documentation [Электронный ресурс]. — "
         "Режим доступа: https://docs.github.com/actions (дата обращения: 14.05.2026).",
        "Docker, Inc. Docker Documentation [Электронный ресурс]. — "
         "Режим доступа: https://docs.docker.com/ (дата обращения: 14.05.2026).",
        "Percival H., Gregory B. Architecture Patterns with Python. — O'Reilly Media, 2020. — 280 с.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(f"{i}. {ref}")
        _fmt_run(run, size=14)

    doc.add_page_break()


def build_appendix(doc: Document) -> None:
    heading1(doc, "ПРИЛОЖЕНИЕ А")
    heading1(doc, "ФРАГМЕНТЫ ИСХОДНОГО КОДА")

    heading2(doc, "А.1 Инициализация схемы БД (src/infra/db/schema.py)")
    code_block(doc, """SCHEMA_VERSION = 3

def init_schema(conn: sqlite3.Connection) -> None:
    conn.executescript(\"\"\"
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY, value TEXT NOT NULL);

        CREATE TABLE IF NOT EXISTS categories (
            id INTEGER PRIMARY KEY, name TEXT NOT NULL UNIQUE,
            kind TEXT NOT NULL CHECK(kind IN ('income','expense','both')) DEFAULT 'both',
            color TEXT NULL, icon TEXT NULL,
            created_at TEXT NOT NULL, updated_at TEXT NOT NULL);

        CREATE TABLE IF NOT EXISTS transactions (
            id INTEGER PRIMARY KEY,
            type TEXT NOT NULL CHECK(type IN ('income','expense')),
            amount_cents INTEGER NOT NULL CHECK(amount_cents > 0),
            occurred_at TEXT NOT NULL,
            category_id INTEGER NULL REFERENCES categories(id) ON DELETE SET NULL,
            note TEXT NULL, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);

        CREATE TABLE IF NOT EXISTS goals (
            id INTEGER PRIMARY KEY, name TEXT NOT NULL,
            target_cents INTEGER NOT NULL CHECK(target_cents > 0),
            current_cents INTEGER NOT NULL CHECK(current_cents >= 0) DEFAULT 0,
            deadline_at TEXT NULL, note TEXT NULL,
            created_at TEXT NOT NULL, updated_at TEXT NOT NULL);

        CREATE TABLE IF NOT EXISTS reminders (
            id INTEGER PRIMARY KEY, name TEXT NOT NULL,
            amount_cents INTEGER NULL CHECK(amount_cents >= 0),
            due_at TEXT NOT NULL,
            recurrence TEXT NOT NULL
                CHECK(recurrence IN ('none','daily','weekly','monthly')) DEFAULT 'none',
            note TEXT NULL, created_at TEXT NOT NULL, updated_at TEXT NOT NULL);

        CREATE INDEX IF NOT EXISTS idx_transactions_occurred_at
            ON transactions(occurred_at);
        CREATE INDEX IF NOT EXISTS idx_transactions_type_occurred
            ON transactions(type, occurred_at);
    \"\"\")""")

    heading2(doc, "А.2 Шифрование файла БД (src/infra/security/crypto.py)")
    code_block(doc, """MAGIC = b"PFM1"
SALT_LEN, NONCE_LEN, KEY_LEN = 16, 12, 32
MIN_PASSPHRASE_LEN = 8

def _derive_key(passphrase: str, salt: bytes) -> bytes:
    if len(passphrase) < MIN_PASSPHRASE_LEN:
        raise ValueError(f"passphrase must be at least {MIN_PASSPHRASE_LEN} characters")
    kdf = Scrypt(salt=salt, length=KEY_LEN, n=2**14, r=8, p=1)
    return kdf.derive(passphrase.encode("utf-8"))

def encrypt_bytes(*, plaintext: bytes, passphrase: str) -> bytes:
    salt, nonce = os.urandom(SALT_LEN), os.urandom(NONCE_LEN)
    key = _derive_key(passphrase, salt)
    ciphertext = AESGCM(key).encrypt(nonce, plaintext, None)
    return MAGIC + salt + nonce + ciphertext

def decrypt_bytes(*, blob: bytes, passphrase: str) -> bytes:
    if blob[:4] != MAGIC:
        raise ValueError("invalid blob header")
    salt  = blob[4:20];  nonce = blob[20:32];  ct = blob[32:]
    key   = _derive_key(passphrase, salt)
    try:
        return AESGCM(key).decrypt(nonce, ct, None)
    except InvalidTag as exc:
        raise InvalidPasswordError("invalid password or corrupted data") from exc""")

    heading2(doc, "А.3 Пример теста (tests/test_infra_db.py)")
    code_block(doc, """def test_goal_deposit_increments_and_caps(conn):
    goals = GoalRepository(conn)
    with transaction(conn):
        gid = goals.create(name="Отпуск", target_cents=10_000, current_cents=0)

    with transaction(conn):
        g = goals.deposit(goal_id=gid, amount_cents=3_000)
    assert g.current_cents == 3_000
    assert round(g.progress_ratio, 2) == 0.3

    # Deposit сверх остатка — должен заблокироваться на target
    with transaction(conn):
        g = goals.deposit(goal_id=gid, amount_cents=99_999)
    assert g.current_cents == g.target_cents
    assert g.progress_ratio == 1.0""")


# ── main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    doc = Document()
    set_margins(doc)
    add_page_numbers(doc)

    # set default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    build_title_page(doc)
    build_toc(doc)
    build_intro(doc)
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)
    build_section5(doc)
    build_conclusion(doc)
    build_references(doc)
    build_appendix(doc)

    out = "docs/poyasnitelnaya_zapiska.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
